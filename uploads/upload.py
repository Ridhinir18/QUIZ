import io
import csv
import re
import fitz
import unicodedata
from docx import Document


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Advanced layout-aware extractor handling multi-column text structures, 
    nested data tables, and multi-language scripts (like Hindi) across multiple file formats.
    Upgraded with PyMuPDF for lightning-fast, column-accurate PDF processing.
    """
    ext = filename.split(".")[-1].lower() if "." in filename else ""
    extracted_text = ""

    if ext == "pdf":
        if fitz is None:
            raise ImportError("The 'pymupdf' library is missing. Please run 'pip install pymupdf'.")
            
        pages_text = []
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            for page in doc:
                text_blocks = page.get_text("blocks")
                
                clean_blocks = [b[4].strip() for b in text_blocks if b[6] == 0 and b[4].strip()]
                
                if clean_blocks:
                    page_content = "\n\n".join(clean_blocks)
                    pages_text.append(page_content)
                    
            doc.close()
            extracted_text = "\n\n".join(pages_text)
        except Exception as e:
            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                extracted_text = "\n\n".join([page.get_text("text") for page in doc])
                doc.close()
            except Exception:
                extracted_text = ""

    elif ext == "docx":
        doc = Document(io.BytesIO(file_bytes))
        elements = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                elements.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                clean_row = []
                for val in row_text:
                    if not clean_row or val != clean_row[-1]:
                        clean_row.append(val)
                if clean_row:
                    elements.append(" | ".join(clean_row))
        extracted_text = "\n\n".join(elements)

    elif ext == "txt":
        try:
            extracted_text = file_bytes.decode("utf-8-sig")
        except UnicodeDecodeError:
            extracted_text = file_bytes.decode("latin-1", errors="ignore")

    elif ext == "csv":
        try:
            decoded = file_bytes.decode("utf-8-sig")
        except UnicodeDecodeError:
            decoded = file_bytes.decode("latin-1", errors="ignore")
        text_stream = io.StringIO(decoded)
        reader = csv.reader(text_stream)
        lines = [" | ".join([col.strip() for col in row if col.strip()]) for row in reader]
        extracted_text = "\n".join([l for l in lines if l])

    if extracted_text:
        extracted_text = unicodedata.normalize("NFKC", extracted_text)
        extracted_text = re.sub(r'\r\n', '\n', extracted_text)
        extracted_text = re.sub(r'[ \t]{2,}', ' ', extracted_text)
        extracted_text = re.sub(r'\n{3,}', '\n\n', extracted_text)
        return extracted_text.strip()
        
    return ""